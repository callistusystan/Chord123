from docx import Document
from docx.shared import RGBColor, Pt
from .translator import Translator

class DocumentHandler:
    def __init__(self, path):
        with open(path, 'rb') as f:
            self.document = Document(f)
            self.content = self.getText(self.document)

    def getText(self, document):
        content = []
        for paragraph in document.paragraphs:
            content.append(paragraph.text)
        return '\n'.join(content)

    def save(self, content, filename):
        document = Document()
        endOfLine1 = content.index('\n')

        # get heading and body
        heading = content[:endOfLine1]
        body = content[endOfLine1+1:]
        lines = body.split('\n')
        while lines[0] == '':
            lines.pop(0)

        # start writing
        document.add_heading(heading, 0)
        t = Translator()
        p = document.add_paragraph()
        for line in lines:
            if line == '':
                p = document.add_paragraph()
            else:
                run = p.add_run(line+'\n')
                run.font.size = Pt(14)
                if t.hasChords(line):
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0xFF)

        document.save(filename)
